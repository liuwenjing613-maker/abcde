#!/usr/bin/env python3
"""Generate Residual v2 architecture Word document."""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_para(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)


def add_bullet(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Bullet")


def add_code_block(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def build_doc(output_path: Path) -> None:
    doc = Document()

    title = doc.add_heading("CCAC_MAPS Residual v2 模型架构说明", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_para(doc, "项目名称：CCAC 青少年心理健康纵向预测（txy 工作区）")
    add_para(doc, "模型名称：ResidualFusionModel + StageWiseLongitudinalModel（Residual v2）")
    add_para(doc, "任务：利用 T1/T2/T3 历史多模态特征与 DASS 量表，预测 T4 焦虑等级（5 分类）")
    add_para(doc, "文档用途：模型结构说明，供团队评审与复现参考")

    add_heading(doc, "一、设计思想（AdoDAS → CCAC 迁移）", 1)
    add_para(
        doc,
        "本模型将 AdoDAS A1 中验证过的「稳定 anchor + 小比例多模态残差修正」思路迁移到 CCAC 纵向五分类任务：",
    )
    add_bullet(doc, "participant-level 聚合：每名被试一行，不做 clip/session 级提交")
    add_bullet(doc, "school/class group split：按 anon_school + anon_class 分组交叉验证，防止班级泄漏")
    add_bullet(doc, "历史量表作 anchor：T1/T2/T3 的 D/A/S 分数与等级及衍生特征作为主信号")
    add_bullet(doc, "多模态作 residual：音视频 SSL 特征以小权重修正 tabular 预测")
    add_bullet(doc, "clip-stage 两层聚合：阶段内 4 clips attention + 跨 T1/T2/T3 时序编码")

    add_heading(doc, "二、总体公式", 1)
    add_code_block(
        doc,
        "logits_final = logits_tabular + α × logits_multimodal\n\n"
        "其中 α = 0.25（默认值）\n"
        "logits_final 经 softmax 得到 5 类概率，argmax 得到焦虑等级预测",
    )

    add_heading(doc, "三、输入数据规格", 1)
    add_para(doc, "每名被试（participant）输入如下：", bold=True)

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "字段"
    hdr[1].text = "张量形状"
    hdr[2].text = "来源"
    hdr[3].text = "说明"

    rows = [
        ("audio", "[B, 3, 4, 3072]", "audio_wavlm_base", "T1/T2/T3 × A01/B01/B02/B03 的 pooled 向量"),
        ("video", "[B, 3, 4, 768]", "video_dinov2_small", "同上，视频 SSL 特征"),
        ("clip_mask", "[B, 3, 4]", "数据完整性", "缺失 clip 置 0，参与 attention 掩码"),
        ("tabular_features", "[B, 39]", "labels.csv 历史量表", "T1-T3 分数/等级 + 斜率/均值等工程特征"),
        ("history_levels", "[B, 9]", "labels.csv 历史等级", "9 个等级槽，供 embedding 使用"),
    ]
    for field, shape, source, note in rows:
        row = table.add_row().cells
        row[0].text = field
        row[1].text = shape
        row[2].text = source
        row[3].text = note

    add_para(doc, "B = batch size；3 = T1/T2/T3 三个阶段；4 = A01/B01/B02/B03 四个 clip。")

    add_heading(doc, "四、模型整体结构图（文字版）", 1)
    add_code_block(
        doc,
        "Person（被试）\n"
        "│\n"
        "├─ Tabular Head（anchor 主分支）\n"
        "│    历史量表 39维 → LayerNorm → Linear(128) → GELU → Dropout → Linear(5)\n"
        "│    输出：logits_tabular [B, 5]\n"
        "│\n"
        "└─ StageWiseLongitudinal（multimodal 残差分支，权重 ×0.25）\n"
        "     │\n"
        "     ├─ 逐 Clip 编码\n"
        "     │    AudioAdapter: 3072 → 256 (Linear+LayerNorm+GELU+Dropout)\n"
        "     │    VideoAdapter:   768 → 256\n"
        "     │    GatedFusion: 门控融合 audio/video → clip 表征 256d\n"
        "     │\n"
        "     ├─ 逐 Stage 聚合（T1/T2/T3 各 4 clips）\n"
        "     │    ClipAttention: 对 4 个 clip 做 softmax 加权池化\n"
        "     │    + 可学习 stage_position 编码\n"
        "     │    输出：stage_repr [B, 3, 256]\n"
        "     │\n"
        "     ├─ 跨 Stage 时序编码\n"
        "     │    BiGRU(256 → 192×2)，建模 T1→T2→T3 轨迹\n"
        "     │    拼接：mean(GRU输出) + mean(stage) + T3最后阶段 → 1152d\n"
        "     │\n"
        "     ├─ HistoryEncoder（多模态侧再用历史量表）\n"
        "     │    分数 MLP(39→128) + 等级 Embedding(9×16→64) → 融合 128d\n"
        "     │\n"
        "     └─ Classifier\n"
        "          concat(1152 + 128) → LayerNorm → Linear(256) → GELU → Linear(5)\n"
        "          输出：logits_multimodal [B, 5]\n"
        "\n"
        "最终：logits_final = logits_tabular + 0.25 × logits_multimodal",
    )

    add_heading(doc, "五、核心模块说明", 1)

    add_heading(doc, "5.1 Tabular Head（历史量表 anchor）", 2)
    add_bullet(doc, "输入 39 维 tabular 特征，包含 T1/T2/T3 的 depression/anxiety/stress 分数与等级")
    add_bullet(doc, "衍生特征：mean/max/slope、T3-T1 差分、是否连续升高等")
    add_bullet(doc, "在训练集上该路信号极强（HistoryTabular OOF macro-F1 ≈ 0.36）")
    add_bullet(doc, "测试集若无历史量表，该分支输入为零，主要依赖多模态分支")

    add_heading(doc, "5.2 FeatureAdapter + GatedFusion", 2)
    add_bullet(doc, "AudioAdapter / VideoAdapter：将不同维度 SSL 特征映射到统一 hidden_dim=256")
    add_bullet(doc, "GatedFusion：学习门控系数 g，输出 g·audio + (1-g)·video，自适应音视频权重")

    add_heading(doc, "5.3 ClipAttentionPool", 2)
    add_bullet(doc, "每个阶段内对 4 个 clip 计算 attention score")
    add_bullet(doc, "对缺失 clip 做 mask（-inf），再 softmax 加权求和")
    add_bullet(doc, "对应 AdoDAS 中 session attention 的 CCAC 版本（session → clip）")

    add_heading(doc, "5.4 BiGRU 时序编码器", 2)
    add_bullet(doc, "输入 T1/T2/T3 三个 stage embedding，输出长度 3 的时序序列")
    add_bullet(doc, "双向 GRU，hidden=192，捕捉焦虑状态纵向演变")
    add_bullet(doc, "融合三种池化：时序均值、stage 均值、最后阶段 T3（强调近期状态）")

    add_heading(doc, "5.5 HistoryEncoder", 2)
    add_bullet(doc, "将历史分数与等级嵌入为 128 维向量")
    add_bullet(doc, "与多模态时序表征拼接后送入分类头")
    add_bullet(doc, "实现「量表 + 音视频」在同一多模态分支内的二次融合")

    add_heading(doc, "5.6 ResidualFusionModel", 2)
    add_bullet(doc, "tabular 与 multimodal 各输出 5 维 logits，线性相加")
    add_bullet(doc, "α=0.25 表示多模态只做小比例修正，避免冲掉量表主信号")
    add_bullet(doc, "思想来源：AdoDAS 后期 anchor + motion signal 小比例融合")

    add_heading(doc, "六、输出与标签编码", 1)
    add_para(doc, "模型内部训练使用按中文排序的类别索引；CodaBench 提交需映射为 DASS 序数：", bold=True)

    table2 = doc.add_table(rows=1, cols=3)
    table2.style = "Table Grid"
    h2 = table2.rows[0].cells
    h2[0].text = "提交 label"
    h2[1].text = "焦虑等级"
    h2[2].text = "分数区间"

    label_rows = [
        ("0", "正常", "0–7"),
        ("1", "轻度", "8–9"),
        ("2", "中度", "10–14"),
        ("3", "重度", "15–19"),
        ("4", "非常严重", "20+"),
    ]
    for a, b, c in label_rows:
        r = table2.add_row().cells
        r[0].text = a
        r[1].text = b
        r[2].text = c

    add_heading(doc, "七、Residual v2 训练配置", 1)
    table3 = doc.add_table(rows=1, cols=2)
    table3.style = "Table Grid"
    table3.rows[0].cells[0].text = "超参数"
    table3.rows[0].cells[1].text = "取值"

    config_rows = [
        ("音频特征", "audio_wavlm_base (3072d)"),
        ("视频特征", "video_dinov2_small (768d)"),
        ("α (residual 权重)", "0.25"),
        ("hidden_dim", "256"),
        ("temporal_hidden_dim (GRU)", "192"),
        ("history_hidden_dim", "128"),
        ("dropout", "0.2"),
        ("optimizer", "AdamW, lr=1e-3, weight_decay=1e-4"),
        ("loss", "CrossEntropyLoss + class_weight^1.5"),
        ("交叉验证", "5-fold GroupKFold (school/class)"),
        ("早停指标", "macro-F1"),
        ("calibrate_bias", "false（不做 val grid-search bias）"),
        ("数据增广", "stage_drop=0.1, clip_drop=0.05, feature_noise=0.01"),
        ("集成方式", "5 折 checkpoint 概率平均"),
    ]
    for k, v in config_rows:
        row = table3.add_row().cells
        row[0].text = k
        row[1].text = v

    add_heading(doc, "八、验证集表现（Residual v2）", 1)
    add_bullet(doc, "OOF macro-F1：0.296")
    add_bullet(doc, "OOF balanced accuracy：0.303")
    add_bullet(doc, "OOF 预测分布：五类均有覆盖（非全塌缩）")
    add_bullet(doc, "测试集预测分布：高度偏向「正常」（380/382），test 泛化仍待改进")

    add_heading(doc, "九、代码与产物路径", 1)
    add_bullet(doc, "核心代码：CCAC/txy/src/txy/models/residual.py, stagewise.py")
    add_bullet(doc, "训练脚本：CCAC/txy/scripts/train_residual.py")
    add_bullet(doc, "配置文件：CCAC/txy/configs/residual_v2.yaml")
    add_bullet(doc, "训练产物：CCAC/txy/artifacts/residual_v2/")
    add_bullet(doc, "提交打包：CCAC/txy/scripts/make_submission.py")

    add_heading(doc, "十、与官方 Baseline 的差异", 1)
    table4 = doc.add_table(rows=1, cols=3)
    table4.style = "Table Grid"
    table4.rows[0].cells[0].text = "维度"
    table4.rows[0].cells[1].text = "官方 Baseline"
    table4.rows[0].cells[2].text = "Residual v2（本模型）"

    diff_rows = [
        ("历史量表", "不使用", "Tabular anchor + HistoryEncoder"),
        ("融合方式", "仅多模态", "tabular + α×multimodal residual"),
        ("CV 策略", "StratifiedKFold", "GroupKFold (school/class)"),
        ("结构", "clip encoder + GRU", "Adapter + GatedFusion + ClipAttn + GRU"),
        ("公开 test 分数", "≈0.23（正确编码后）", "分布过保守，待优化"),
    ]
    for a, b, c in diff_rows:
        r = table4.add_row().cells
        r[0].text = a
        r[1].text = b
        r[2].text = c

    add_para(doc, "")
    add_para(doc, "—— 文档完 ——", bold=True)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


if __name__ == "__main__":
    out = Path("/home/adodas/CCAC/txy/docs/Residual_v2_模型架构说明.docx")
    build_doc(out)
    print(f"saved: {out}")
